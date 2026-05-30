from __future__ import annotations

from importlib.util import module_from_spec, spec_from_file_location
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_PATH = ROOT / "docs" / "a3_python_course_framework.docx"
FALLBACK_OUT_PATH = ROOT / "docs" / "a3_python_project_implementation_framework.docx"
TABLE_GEOMETRY_PATH = Path(
    r"C:\Users\Lenovo\.codex\plugins\cache\openai-primary-runtime\documents\26.515.10909\skills\documents\scripts\table_geometry.py"
)


def load_table_geometry():
    spec = spec_from_file_location("table_geometry", TABLE_GEOMETRY_PATH)
    if spec is None or spec.loader is None:
        raise RuntimeError(f"Unable to load table geometry helper: {TABLE_GEOMETRY_PATH}")
    module = module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


table_geometry = load_table_geometry()
apply_table_geometry = table_geometry.apply_table_geometry


BODY_FONT = "Calibri"
EAST_ASIA_FONT = "Microsoft YaHei"
TEXT_COLOR = RGBColor.from_string("000000")
MUTED_COLOR = RGBColor.from_string("666666")
HEADING_BLUE = RGBColor.from_string("2E74B5")
HEADING_DARK = RGBColor.from_string("1F4D78")
TABLE_HEADER_FILL = "F2F4F7"
TABLE_SOFT_FILL = "E8EEF5"


def set_run_font(run, *, size=11, bold=False, color=TEXT_COLOR):
    run.font.name = BODY_FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), BODY_FONT)
    rfonts.set(qn("w:hAnsi"), BODY_FONT)
    rfonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    rfonts.set(qn("w:cs"), BODY_FONT)


def set_para(paragraph, *, before=0, after=6, line_spacing=1.10, align=None):
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line_spacing
    fmt.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    if align is not None:
        paragraph.alignment = align


def add_p(doc: Document, text: str, *, size=11, bold=False, color=TEXT_COLOR, align=None):
    paragraph = doc.add_paragraph()
    set_para(paragraph, align=align)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    return paragraph


def add_bullet(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Bullet")
    set_para(paragraph, before=0, after=3)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=TEXT_COLOR)
    return paragraph


def add_number(doc: Document, text: str):
    paragraph = doc.add_paragraph(style="List Number")
    set_para(paragraph, before=0, after=4)
    run = paragraph.add_run(text)
    set_run_font(run, size=11, color=TEXT_COLOR)
    return paragraph


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell(cell, text: str, *, bold=False, fill=None, align=None, size=10.5):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.LEFT
    set_para(paragraph, before=0, after=0, line_spacing=1.0)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold, color=TEXT_COLOR)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        shade_cell(cell, fill)


def add_table(doc: Document, rows: list[tuple[str, ...]], widths: list[int], *, header=True):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row_idx, row_values in enumerate(rows):
        for col_idx, value in enumerate(row_values):
            set_cell(
                table.rows[row_idx].cells[col_idx],
                value,
                bold=header and row_idx == 0,
                fill=TABLE_HEADER_FILL if header and row_idx == 0 else None,
                align=WD_ALIGN_PARAGRAPH.CENTER if header and row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT,
            )
    apply_table_geometry(table, widths)
    doc.add_paragraph()
    return table


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = TEXT_COLOR
    normal._element.rPr.rFonts.set(qn("w:ascii"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), BODY_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    normal._element.rPr.rFonts.set(qn("w:cs"), BODY_FONT)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

    title = doc.styles["Title"]
    title.font.name = BODY_FONT
    title.font.size = Pt(20)
    title.font.bold = True
    title.font.color.rgb = TEXT_COLOR
    title._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)

    for style_name, size, color, before, after in [
        ("Heading 1", 16, HEADING_BLUE, 16, 8),
        ("Heading 2", 13, HEADING_BLUE, 12, 6),
        ("Heading 3", 12, HEADING_DARK, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = BODY_FONT
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style._element.rPr.rFonts.set(qn("w:eastAsia"), EAST_ASIA_FONT)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.10
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE


def build_document():
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("智学工坊项目实施框架")
    set_run_font(run, size=20, bold=True)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_para(subtitle, before=0, after=10, line_spacing=1.0)
    run = subtitle.add_run("基于大模型的《Python程序设计》个性化资源生成与学习多智能体系统")
    set_run_font(run, size=11, color=MUTED_COLOR)

    doc.add_heading("1. 项目目标", level=1)
    add_p(
        doc,
        "本项目建设一个面向《Python程序设计》课程的个性化学习系统。系统通过对话收集学生学习情况，形成动态画像，再由多个智能体协同生成学习路径、课程讲义、练习题、代码案例、错题分析和复习计划，最终通过测验结果反向更新学生画像，实现“画像-资源-学习-评估-再推荐”的完整闭环。"
    )
    add_bullet(doc, "首版只服务一门课程：《Python程序设计》。")
    add_bullet(doc, "首版只保证一个核心场景：学生自主学习 Python 并获得个性化资源。")
    add_bullet(doc, "所有 AI 输出必须结构化、可落库、可展示、可追踪。")
    add_bullet(doc, "系统不是普通聊天框，必须展示多智能体协作过程和学习闭环。")

    doc.add_heading("2. 首版可运行闭环", level=1)
    flow_rows = [
        ("步骤", "输入", "系统处理", "输出"),
        ("1. 建立画像", "专业、基础、目标、薄弱点、学习偏好、可投入时间", "画像智能体抽取特征并打标签", "学生画像 JSON、薄弱点列表"),
        ("2. 规划路径", "学生画像、课程知识节点、当前进度", "路径规划智能体生成阶段任务", "个性化学习路径、推荐理由"),
        ("3. 生成资源", "路径节点、学生画像、课程资料", "资源生成智能体和题目智能体协同生成内容", "讲义、例题、代码案例、练习题、复习卡片"),
        ("4. 智能答疑", "学生问题、当前知识节点、历史学习数据", "问答智能体结合知识库回答", "Markdown 答案、代码示例、延伸建议"),
        ("5. 测验反馈", "练习答案、学习行为、错题", "评估智能体分析掌握程度", "得分、错题原因、画像更新、下一步推荐"),
    ]
    add_table(doc, flow_rows, [850, 2450, 3150, 2910])

    doc.add_heading("3. 课程内容范围", level=1)
    course_rows = [
        ("模块", "知识点", "系统内资源", "验收标准"),
        ("Python 基础", "变量、数据类型、输入输出、运算符", "概念卡片、基础例题、随堂练习", "能为零基础学生生成入门学习包"),
        ("流程控制", "if 分支、for/while 循环、break/continue", "流程图讲解、易错题、调试案例", "能解释循环边界和条件判断错误"),
        ("函数与模块", "函数定义、参数、返回值、模块导入", "函数拆解图、代码重构建议", "能生成函数化改写案例"),
        ("组合数据类型", "列表、元组、字典、集合、字符串处理", "对比表、操作速查、应用练习", "能按学生薄弱点推送专项题"),
        ("文件与异常", "文件读写、异常捕获、上下文管理", "实验文档、代码模板、错误解释", "能生成可运行的小实验"),
        ("面向对象与综合项目", "类、对象、封装、简单项目组织", "项目案例、分步任务、评分标准", "能形成课程综合实践任务"),
    ]
    add_table(doc, course_rows, [1300, 2450, 2800, 2810])

    doc.add_heading("4. 功能模块设计", level=1)
    module_rows = [
        ("模块", "必须实现的功能", "关键数据", "前端页面"),
        ("学生画像", "对话式信息采集、标签抽取、画像更新、薄弱点识别", "profile、tags、weaknesses、preferences", "/student/profile"),
        ("课程知识库", "知识节点维护、资料片段检索、引用来源记录", "course_units、knowledge_items", "/admin/course-units"),
        ("学习路径", "按画像生成阶段路径、任务顺序、预计时长和推荐理由", "learning_paths、path_steps", "/student/learning-path"),
        ("资源生成", "生成讲义、练习、代码案例、复习卡片、思维导图", "resources、resource_items", "/student/resources"),
        ("智能答疑", "课程内问答、代码解释、错题讲解、继续追问", "chat_sessions、chat_messages", "/student/tutor"),
        ("测验评估", "自动组卷、提交评分、错题分析、掌握度更新", "quizzes、quiz_attempts、mastery_records", "/student/quiz"),
        ("安全审核", "事实性检查、敏感内容过滤、输出格式校验", "agent_runs、safety_flags", "后台日志页面"),
    ]
    add_table(doc, module_rows, [1300, 3400, 2450, 2210])

    doc.add_heading("5. 多智能体设计", level=1)
    agent_rows = [
        ("智能体", "触发时机", "输入", "输出"),
        ("ProfileAgent 画像智能体", "学生提交画像问卷或完成测验后", "学生自然语言描述、测验结果、历史记录", "画像标签、基础等级、薄弱点、学习偏好"),
        ("PlannerAgent 路径规划智能体", "画像生成后或画像更新后", "画像、课程节点、掌握度", "学习阶段、节点顺序、预计时长、推荐理由"),
        ("ResourceAgent 资源生成智能体", "学生选择路径节点后", "知识节点、学生画像、课程资料片段", "讲义、知识卡片、代码案例、拓展材料"),
        ("ExerciseAgent 练习生成智能体", "生成资源包或复习包时", "知识节点、难度、薄弱点", "选择题、填空题、代码题、答案解析"),
        ("TutorAgent 答疑智能体", "学生提问时", "问题、当前节点、知识库检索结果", "解释、示例、代码、后续建议"),
        ("EvaluatorAgent 评估智能体", "测验提交后", "答案、得分、错题、学习记录", "掌握度、错因、补救路径"),
        ("SafetyAgent 审核智能体", "所有 AI 内容返回前", "待输出内容、引用资料、规则", "通过/驳回/重写建议"),
    ]
    add_table(doc, agent_rows, [1700, 2100, 2800, 2760])

    doc.add_heading("6. 数据库表结构草案", level=1)
    db_rows = [
        ("表名", "核心字段", "说明"),
        ("users", "id, username, password_hash, role, created_at", "用户基础表，首版可保留 student/admin 两类"),
        ("student_profiles", "id, user_id, major, goal, base_level, preferences, weaknesses_json, tags_json, updated_at", "学生画像主表"),
        ("course_units", "id, title, order_index, description, difficulty, prerequisites_json", "Python 课程知识节点"),
        ("knowledge_items", "id, unit_id, title, content, source_type, source_ref, keywords_json", "知识库资料片段"),
        ("learning_paths", "id, user_id, profile_version, status, created_at", "某次生成的个性化路径"),
        ("path_steps", "id, path_id, unit_id, step_order, goal, estimated_minutes, reason", "路径中的阶段任务"),
        ("resources", "id, user_id, unit_id, type, title, content_markdown, metadata_json, created_at", "AI 生成资源"),
        ("quizzes", "id, unit_id, title, difficulty, questions_json", "测验与题目集合"),
        ("quiz_attempts", "id, user_id, quiz_id, answers_json, score, analysis_json, created_at", "学生测验记录"),
        ("mastery_records", "id, user_id, unit_id, mastery_score, weakness_tags_json, updated_at", "知识点掌握度"),
        ("agent_runs", "id, user_id, agent_name, input_json, output_json, status, latency_ms, created_at", "智能体运行日志"),
        ("chat_messages", "id, user_id, session_id, role, content, related_unit_id, created_at", "问答历史"),
    ]
    add_table(doc, db_rows, [1900, 5050, 2410])

    doc.add_heading("7. 后端接口设计", level=1)
    api_rows = [
        ("接口", "请求体关键字段", "响应体关键字段", "用途"),
        ("POST /api/profile/analyze", "answers, freeText", "profile, tags, weaknesses", "生成或更新学生画像"),
        ("GET /api/course/units", "无", "units[]", "获取 Python 课程知识节点"),
        ("POST /api/path/generate", "profileId, targetUnitIds", "path, steps[]", "生成学习路径"),
        ("GET /api/path/current", "无", "path, steps, progress", "获取当前路径和进度"),
        ("POST /api/resources/generate", "unitId, resourceTypes[]", "resources[]", "生成个性化资源包"),
        ("GET /api/resources", "unitId, type", "resources[]", "查询已生成资源"),
        ("POST /api/tutor/chat", "message, unitId, sessionId", "reply, citations, suggestions", "课程问答"),
        ("POST /api/quiz/generate", "unitId, difficulty, count", "quiz", "生成测验"),
        ("POST /api/quiz/submit", "quizId, answers", "score, analysis, profilePatch", "评分并更新画像"),
        ("GET /api/progress/overview", "无", "mastery, recentActivities, nextActions", "学习进度总览"),
    ]
    add_table(doc, api_rows, [2200, 2550, 2550, 2060])

    doc.add_heading("8. AI 输出结构约束", level=1)
    add_p(doc, "所有智能体不直接返回散文式长文本，后端必须要求模型输出 JSON，再由前端按类型渲染。以下为核心资源包的建议结构：")
    contract_rows = [
        ("字段", "类型", "说明"),
        ("title", "string", "资源标题，例如“循环语句专项学习包”"),
        ("unitId", "string", "关联课程知识节点"),
        ("profileReason", "string", "为什么适合该学生"),
        ("lectureMarkdown", "string", "讲义内容，支持 Markdown"),
        ("codeExamples", "array", "代码案例，包含 title、code、explanation"),
        ("exercises", "array", "练习题，包含 type、question、options、answer、explanation"),
        ("mindmapMermaid", "string", "可选，知识结构图 Mermaid 文本"),
        ("reviewCards", "array", "复习卡片，包含 question、answer"),
        ("safetyStatus", "string", "passed / rewritten / blocked"),
    ]
    add_table(doc, contract_rows, [2100, 1300, 5960])

    doc.add_heading("9. 前端页面实施清单", level=1)
    page_rows = [
        ("页面", "核心组件", "需要调用的接口"),
        ("/student/profile", "画像问卷、对话输入、标签展示", "POST /api/profile/analyze"),
        ("/student/home", "学习概览、当前任务、最近资源", "GET /api/progress/overview"),
        ("/student/learning-path", "路径时间线、推荐理由、阶段状态", "POST /api/path/generate, GET /api/path/current"),
        ("/student/resources", "资源卡片、Markdown 渲染、代码块、思维导图", "POST /api/resources/generate, GET /api/resources"),
        ("/student/tutor", "聊天窗口、引用来源、追问建议", "POST /api/tutor/chat"),
        ("/student/quiz", "答题区、提交、错题解析", "POST /api/quiz/generate, POST /api/quiz/submit"),
        ("/admin/course-units", "课程节点管理、资料录入、题库管理", "GET /api/course/units"),
    ]
    add_table(doc, page_rows, [2200, 3900, 3260])

    doc.add_heading("10. 课程 seed 数据要求", level=1)
    add_bullet(doc, "至少录入 6 个课程模块，对应第 3 节中的 Python 知识范围。")
    add_bullet(doc, "每个模块至少包含 5 条知识资料片段、5 道基础题、3 道进阶题、1 个代码案例。")
    add_bullet(doc, "每条资料片段必须记录 source_ref，后续展示为“内容依据”。")
    add_bullet(doc, "每个模块至少准备一个常见错误，例如变量命名、循环边界、列表索引、函数返回值、文件路径错误。")
    add_bullet(doc, "演示前至少准备 3 个典型学生画像：零基础型、应试补弱型、项目实践型。")

    doc.add_heading("11. 实施顺序", level=1)
    roadmap_rows = [
        ("阶段", "实现内容", "完成标志"),
        ("阶段 1：数据与基础接口", "建立 course_units、knowledge_items、student_profiles；完成课程节点查询和画像生成接口", "能提交画像并看到结构化画像结果"),
        ("阶段 2：路径与资源", "实现 PlannerAgent、ResourceAgent、ExerciseAgent；完成路径生成和资源生成接口", "能为一个学生生成完整学习包"),
        ("阶段 3：问答与测验", "实现 TutorAgent、EvaluatorAgent；接入测验生成、提交评分和错因分析", "能完成一次学习-测试-反馈闭环"),
        ("阶段 4：安全与展示", "加入 SafetyAgent、日志、错误处理、加载状态和演示数据", "演示流程稳定，失败时有兜底文案"),
    ]
    add_table(doc, roadmap_rows, [1700, 5300, 2360])

    doc.add_heading("12. 验收清单", level=1)
    checklist = [
        "学生填写画像后，系统能生成不少于 6 个维度的画像标签。",
        "系统能围绕一个 Python 知识节点生成不少于 5 类资源。",
        "学习路径必须包含阶段顺序、预计时间、推荐理由和当前状态。",
        "智能答疑必须能结合课程资料回答，并给出引用或依据说明。",
        "测验提交后必须产生得分、错题原因、补救资源和画像更新。",
        "后台必须记录智能体运行日志，便于展示多智能体协同过程。",
        "所有 AI 内容返回前必须通过格式校验和安全校验。",
        "系统演示必须能在 7 分钟内完整走通一个学生学习闭环。",
    ]
    for item in checklist:
        add_bullet(doc, item)

    doc.add_heading("13. 首版目录落地建议", level=1)
    tree_rows = [
        ("路径", "内容"),
        ("backend/src/routes/profileRoutes.ts", "画像生成与画像查询接口"),
        ("backend/src/routes/pathRoutes.ts", "学习路径生成与查询接口"),
        ("backend/src/routes/resourceRoutes.ts", "学习资源生成与查询接口"),
        ("backend/src/routes/tutorRoutes.ts", "智能答疑接口"),
        ("backend/src/routes/quizRoutes.ts", "测验生成、提交和分析接口"),
        ("backend/src/agents", "ProfileAgent、PlannerAgent、ResourceAgent、ExerciseAgent、TutorAgent、EvaluatorAgent、SafetyAgent"),
        ("backend/src/services", "数据库访问、课程资料检索、资源落库、进度更新"),
        ("frontend/src/pages/student", "画像页、首页、路径页、资源页、答疑页、测验页"),
        ("docs", "接口说明、测试说明、部署说明、演示脚本"),
    ]
    add_table(doc, tree_rows, [3600, 5760])

    closing = doc.add_paragraph()
    set_para(closing, before=0, after=0)
    run = closing.add_run("后续开发应优先保证主闭环可运行，再逐步补充页面美化、教师视图和管理视图。只要主闭环稳定，项目就具备完整演示和继续扩展的基础。")
    set_run_font(run, size=11)

    return doc


def main():
    doc = build_document()
    OUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    try:
        doc.save(OUT_PATH)
        print(f"saved: {OUT_PATH}")
    except PermissionError:
        doc.save(FALLBACK_OUT_PATH)
        print(f"saved: {FALLBACK_OUT_PATH}")


if __name__ == "__main__":
    main()
